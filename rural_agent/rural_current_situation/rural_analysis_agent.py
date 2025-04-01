import os
import json
import asyncio
from typing import Dict, List, Optional, Union, Any
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

# 导入文档处理库
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
# from weasyprint import HTML
import tempfile

# 导入GPTResearcher相关模块
from gpt_researcher.agent import GPTResearcher
from gpt_researcher.utils.enum import ReportSource, ReportType, Tone

class RuralAnalysisAgent:
    """乡村现状分析智能体，基于GPTResearcher进行分析并生成结构化报告"""
    
    def __init__(self, framework_path: str = None):
        """
        初始化乡村分析智能体
        
        Args:
            framework_path: 分析框架JSON文件路径，默认使用预设位置
        """
        # 加载环境变量
        load_dotenv()
        
        # 设置分析框架路径       
        if framework_path is None:
            framework_path = os.path.join(
                os.path.dirname(os.path.abspath(__file__)),
                "framework.json"
            )


        # 加载分析框架
        self.framework = self._load_framework(framework_path)
        
        # 输出目录设置
        self.output_dir = os.path.join(
            os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
            "output"
        )
        os.makedirs(self.output_dir, exist_ok=True)
    
    def _load_framework(self, framework_path: str) -> Dict:
        """
        加载分析框架
        
        Args:
            framework_path: 框架JSON文件路径
            
        Returns:
            加载的框架字典
        """
        try:
            with open(framework_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"加载框架文件失败: {e}")
            return {"乡村现状分析框架": {}} # 返回空框架作为默认值
    
    async def analyze(self, query_config: Union[str, Dict]) -> str:
        """
        根据查询配置执行乡村现状分析
        
        Args:
            query_config: JSON字符串或字典，包含查询配置
            
        Returns:
            生成的分析报告路径
        """
        # 转换查询配置为字典格式
        if isinstance(query_config, str):
            try:
                query_config = json.loads(query_config)
            except json.JSONDecodeError:
                raise ValueError("查询配置JSON格式无效")
        
        # 解析查询配置
        village_name = query_config.get("village_name", "未指定乡村")
        query = query_config.get("query", f"{village_name}乡村现状分析")
        report_source = query_config.get("report_source", "web") # web, docs, hybrid
        report_format = query_config.get("report_format", "markdown") # markdown, docx, pdf
        output_path = query_config.get("output_path", self.output_dir)


        # 确保输出目录存在
        try:
            os.makedirs(output_path, exist_ok=True)
            print(f"输出目录已准备: {output_path}")
        except Exception as e:
            print(f"创建输出目录失败: {e}")
            # 如果无法创建指定目录，回退到默认目录
            output_path = self.output_dir
            print(f"使用默认输出目录: {output_path}")

        # 处理数据源配置
        source_urls = query_config.get("source_urls", None)
        document_urls = query_config.get("document_urls", None)
        complement_source_urls = query_config.get("complement_source_urls", True)
        query_domains = query_config.get("query_domains", None)
        
        # 设置GPTResearcher配置
        if report_source == "web":
            report_source_enum = ReportSource.Web.value
            documents = None
        elif report_source == "docs":
            report_source_enum = ReportSource.Local.value
            documents = self._prepare_documents(document_urls)
        else:  # hybrid
            report_source_enum = ReportSource.Hybrid.value
            documents = self._prepare_documents(document_urls) if document_urls else None
        
        # 修改查询内容，强制使用中文
        chinese_query = f"用中文分析：{query}（必须使用中文回答）"

        # 初始化GPTResearcher
        researcher = GPTResearcher(
            query=chinese_query,
            report_type=ReportType.DetailedReport.value,
            report_format="markdown",  # 始终使用markdown作为中间格式
            report_source=report_source_enum,
            tone=Tone.Objective,
            source_urls=source_urls,
            document_urls=document_urls,
            complement_source_urls=complement_source_urls,
            query_domains=query_domains,
            documents=documents,
            verbose=True
        )
        
        # # 执行研究
        # context = await researcher.conduct_research()
        
        # # 将框架转换为提示指南
        # framework_prompt = self._framework_to_prompt()
        
        # # 向上下文添加框架指导
        # context.append({
        #     "type": "framework",
        #     "content": framework_prompt
        # })


        search_summary = await researcher.conduct_research()

        # 将框架转换为提示指南
        framework_prompt = self._framework_to_prompt()

        # 创建一个新的上下文列表，包含搜索结果和框架
        ext_context = [
            {
                "type": "language_instruction",
                "content": "这是一份中文分析报告，整个报告必须完全使用中文撰写，不允许使用英文。即使搜索到的材料是英文，也请将内容翻译成中文后再整合到报告中。"
            },
            {
                "type": "search_results",
                "content": search_summary  # 搜索结果作为第一个元素
            },
            {
                "type": "framework",
                "content": framework_prompt  # 框架作为第二个元素
            }
        ]
        # print(f"扩展上下文内容: {ext_context}")

        # 使用组合的上下文生成报告
        report_markdown = await researcher.write_report(ext_context=ext_context)
                
        
        # 添加参考文献
        report_markdown = researcher.add_references(report_markdown, researcher.visited_urls)
        
        # 根据指定格式输出报告
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename_base = f"{village_name}_乡村现状分析_{timestamp}"
        
        if report_format == "markdown":
            output_file = os.path.join(output_path, f"{filename_base}.md")
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(report_markdown)
        elif report_format == "docx":
            output_file = os.path.join(output_path, f"{filename_base}.docx")
            self._markdown_to_docx(report_markdown, output_file, village_name)
        elif report_format == "pdf":
            output_file = os.path.join(output_path, f"{filename_base}.pdf")
            self._markdown_to_pdf(report_markdown, output_file)
        else:
            # 默认输出markdown
            output_file = os.path.join(output_path, f"{filename_base}.md")
            with open(output_file, "w", encoding="utf-8") as f:
                f.write(report_markdown)
        
        return output_file
    
    def _prepare_documents(self, document_urls: List[str]) -> List[Dict]:
        """处理文档URL，返回文档内容列表"""
        # 这里可以实现文档加载逻辑，如果需要的话
        # 当前版本简单返回None，实际使用时需要完善
        return None
    
    def _framework_to_prompt(self) -> str:
        """将框架转换为LLM提示格式"""
        prompt = "请务必使用中文对乡村现状进行全面分析，按照以下框架结构撰写完整的中文分析报告，确保覆盖所有关键维度：\n\n"
        
        framework = self.framework.get("乡村现状分析框架", {})
        
        for main_category, subcategories in framework.items():
            prompt += f"# {main_category}\n\n"
            
            for subcategory_name, items in subcategories.items():
                prompt += f"## {subcategory_name}\n\n"
                
                if isinstance(items, list):
                    for item in items:
                        prompt += f"- {item}\n"
                elif isinstance(items, dict):
                    for item_name, subitems in items.items():
                        prompt += f"### {item_name}\n\n"
                        for subitem in subitems:
                            prompt += f"- {subitem}\n"
                
                prompt += "\n"
        
        prompt += """
        请确保分析报告结构清晰，内容全面，数据准确，并提供具体、详细的现状描述和数据支持。
        整个报告必须使用中文撰写，除了必要的英文单词、句子和超链接之外，不允许出现英文内容。
        """
        # print(f"生成的提示内容: {prompt}")
        return prompt    

    
    def _markdown_to_docx(self, markdown_text: str, output_path: str, village_name: str) -> None:
        """
        将Markdown转换为Word文档
        
        Args:
            markdown_text: Markdown格式的文本
            output_path: 输出Word文档路径
            village_name: 乡村名称，用于标题
        """
        # 创建新文档
        doc = Document()
        
        # 尝试使用多种字体名称 - 同一个字体的不同表示方式
        chinese_fonts = ['宋体', 'SimSun', '宋体 (简体中文)']
        
        # 设置文档默认样式
        styles = doc.styles
        
        # 设置Normal样式
        normal_style = styles['Normal']
        normal_font = normal_style.font
        # 同时尝试多个字体名称
        for font_name in chinese_fonts:
            normal_font.name = font_name
            if hasattr(normal_font, 'name_ascii'):
                normal_font.name_ascii = font_name
            if hasattr(normal_font, 'name_eastasia'):
                normal_font.name_eastasia = font_name
            if hasattr(normal_font, 'name_hAnsi'):
                normal_font.name_hAnsi = font_name
            if hasattr(normal_font, 'name_cs'):
                normal_font.name_cs = font_name
        
        normal_font.size = Pt(12)  # 小四号字=12磅
        
        # 设置所有标题样式
        for i in range(1, 10):
            if f'Heading {i}' in styles:
                heading_style = styles[f'Heading {i}']
                if hasattr(heading_style, 'font'):
                    heading_font = heading_style.font
                    # 同样尝试多个字体名称
                    for font_name in chinese_fonts:
                        heading_font.name = font_name
                        if hasattr(heading_font, 'name_ascii'):
                            heading_font.name_ascii = font_name
                        if hasattr(heading_font, 'name_eastasia'):
                            heading_font.name_eastasia = font_name
                        if hasattr(heading_font, 'name_hAnsi'):
                            heading_font.name_hAnsi = font_name
                        if hasattr(heading_font, 'name_cs'):
                            heading_font.name_cs = font_name
        
        # 辅助函数：设置中文字体 - 尝试多种字体名称
        def set_chinese_font(run, font_size, bold=False, italic=False):
            # 尝试多个字体名称
            for font_name in chinese_fonts:
                run.font.name = font_name
                if hasattr(run.font, 'name_ascii'):
                    run.font.name_ascii = font_name
                if hasattr(run.font, 'name_eastasia'):
                    run.font.name_eastasia = font_name
                if hasattr(run.font, 'name_hAnsi'):
                    run.font.name_hAnsi = font_name
                if hasattr(run.font, 'name_cs'):
                    run.font.name_cs = font_name
            
            run.font.size = Pt(font_size)
            run.font.bold = bold
            run.font.italic = italic
        
        # 处理粗体和其他格式的辅助函数
        def process_markdown_formatting(paragraph, text, font_size, is_reference=False):
            import re
            
            # 先处理粗体文本 (**text** 或 __text__)
            bold_pattern = r'\*\*(.*?)\*\*|__(.*?)__'
            bold_matches = list(re.finditer(bold_pattern, text))
            
            if not bold_matches:
                # 没有粗体标记，直接添加普通文本
                run = paragraph.add_run(text)
                set_chinese_font(run, font_size if not is_reference else 10.5)
                return
            
            # 处理带有粗体的文本
            last_end = 0
            for match in bold_matches:
                # 添加粗体前的普通文本
                if match.start() > last_end:
                    normal_text = text[last_end:match.start()]
                    run = paragraph.add_run(normal_text)
                    set_chinese_font(run, font_size if not is_reference else 10.5)
                
                # 添加粗体文本
                bold_text = match.group(1) if match.group(1) is not None else match.group(2)
                run = paragraph.add_run(bold_text)
                set_chinese_font(run, font_size if not is_reference else 10.5, bold=True)
                
                last_end = match.end()
            
            # 添加最后一个粗体后的普通文本
            if last_end < len(text):
                normal_text = text[last_end:]
                run = paragraph.add_run(normal_text)
                set_chinese_font(run, font_size if not is_reference else 10.5)
        
        # 添加文档标题 (二号字)
        title = doc.add_heading(level=0)
        # title_run = title.add_run(f"{village_name} - 乡村现状分析报告")
        title_run = title.add_run(f"{village_name}\n乡村现状分析报告")
        set_chinese_font(title_run, 22)  # 二号字
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加生成时间
        time_para = doc.add_paragraph()
        time_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        time_run = time_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        set_chinese_font(time_run, 12)  # 小四号字
        
        # Markdown解析并转换为Word文档
        lines = markdown_text.split('\n')
        
        current_list = []
        in_list = False
        in_references = False
        
        # 先遍历标识参考文献部分
        for i, line in enumerate(lines):
            if '参考文献' in line or 'References' in line:
                in_references = True
                break
        
        # 主处理循环
        for line in lines:
            # 更新参考文献标志
            if '参考文献' in line or 'References' in line:
                in_references = True
                
            # 处理标题和内容
            if line.startswith('# '):
                # 处理一级标题 (三号字)
                if in_list:
                    for item in current_list:
                        p = doc.add_paragraph(style='List Bullet')
                        process_markdown_formatting(p, item, 12, in_references)
                    current_list = []
                    in_list = False
                
                heading = doc.add_heading('', level=1)
                run = heading.add_run(line[2:])
                set_chinese_font(run, 16)  # 三号字
                
            elif line.startswith('## '):
                # 处理二级标题 (四号字)
                if in_list:
                    for item in current_list:
                        p = doc.add_paragraph(style='List Bullet')
                        process_markdown_formatting(p, item, 12, in_references)
                    current_list = []
                    in_list = False
                
                heading = doc.add_heading('', level=2)
                run = heading.add_run(line[3:])
                set_chinese_font(run, 14)  # 四号字
                
            elif line.startswith('### '):
                # 处理三级标题 (小四号字)
                if in_list:
                    for item in current_list:
                        p = doc.add_paragraph(style='List Bullet')
                        process_markdown_formatting(p, item, 12, in_references)
                    current_list = []
                    in_list = False
                
                heading = doc.add_heading('', level=3)
                run = heading.add_run(line[4:])
                set_chinese_font(run, 12)  # 小四号字
                
            elif line.startswith('####') or line.startswith('#####') or line.startswith('######'):
                # 处理四级及以下标题 (小四号字，不使用斜体)
                if in_list:
                    for item in current_list:
                        p = doc.add_paragraph(style='List Bullet')
                        process_markdown_formatting(p, item, 12, in_references)
                    current_list = []
                    in_list = False
                
                level = len(line) - len(line.lstrip('#'))
                heading = doc.add_heading('', level=min(level, 9))
                run = heading.add_run(line[level+1:])
                set_chinese_font(run, 12, italic=False)  # 小四号字，明确设置不使用斜体
                
            elif line.startswith('- ') or line.startswith('* '):
                # 处理列表项
                in_list = True
                current_list.append(line[2:])
                
            elif line.strip() == '':
                # 处理空行
                if in_list:
                    for item in current_list:
                        p = doc.add_paragraph(style='List Bullet')
                        process_markdown_formatting(p, item, 12, in_references)
                    current_list = []
                    in_list = False
                
                # 添加空段落
                if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text != '':
                    p = doc.add_paragraph()
                    
            else:
                # 处理普通段落
                if in_list:
                    for item in current_list:
                        p = doc.add_paragraph(style='List Bullet')
                        process_markdown_formatting(p, item, 12, in_references)
                    current_list = []
                    in_list = False
                
                p = doc.add_paragraph()
                process_markdown_formatting(p, line, 12, in_references)
        
        # 处理末尾的列表项
        if in_list:
            for item in current_list:
                p = doc.add_paragraph(style='List Bullet')
                process_markdown_formatting(p, item, 12, in_references)
        
        # 保存文档前进行直接XML修改以确保字体设置
        try:
            from docx.oxml.ns import qn
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:ascii'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:hAnsi'), '宋体')
                    run._element.rPr.rFonts.set(qn('w:cs'), '宋体')
        except:
            print("无法应用高级字体设置，但文档仍将保存")
        
        # 保存文档
        doc.save(output_path)

    
    # def _markdown_to_pdf(self, markdown_text: str, output_path: str) -> None:
    #     """
    #     将Markdown转换为PDF
        
    #     Args:
    #         markdown_text: Markdown格式的文本
    #         output_path: 输出PDF文档路径
    #     """
    #     # 转换markdown为HTML
    #     html_content = markdown.markdown(markdown_text)
        
    #     # 添加基本样式
    #     styled_html = f"""
    #     <html>
    #     <head>
    #         <meta charset="UTF-8">
    #         <style>
    #             body {{ font-family: Arial, sans-serif; margin: 2cm; }}
    #             h1 {{ color: #333366; }}
    #             h2 {{ color: #336699; }}
    #             h3 {{ color: #339999; }}
    #             .date {{ text-align: right; color: gray; }}
    #         </style>
    #     </head>
    #     <body>
    #         <div class="date">生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
    #         {html_content}
    #     </body>
    #     </html>
    #     """
        
    #     # 使用WeasyPrint生成PDF
    #     HTML(string=styled_html).write_pdf(output_path)

# 主函数
async def main():
    # 加载.env文件
    load_dotenv()
    
    # 创建乡村分析智能体
    agent = RuralAnalysisAgent()
    
    # 示例查询配置
    query_config = {
        "village_name": "示例村",
        "query": "示例村的乡村现状分析，包括自然条件、社会经济结构、基础设施等",
        "report_source": "hybrid",  # web, docs, hybrid
        "report_format": "docx",    # markdown, docx, pdf
        "source_urls": [],  # 可以提供特定网页URL
        "document_urls": [],  # 可以提供本地文档路径
        "complement_source_urls": True,  # 是否自动补充搜索结果
        "query_domains": []  # 限定搜索域名
    }
    
    # 执行分析
    result_path = await agent.analyze(query_config)
    print(f"已生成分析报告: {result_path}")

if __name__ == "__main__":
    asyncio.run(main())